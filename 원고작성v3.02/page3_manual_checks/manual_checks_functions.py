import os
from collections import Counter
from konlpy.tag import Okt

def count_characters(text):
    return len(text)

def count_frequent_words(text, min_count=2):
    okt = Okt()
    words = okt.nouns(text)
    word_counts = Counter(words)
    frequent_words = [(word, count) for word, count in word_counts.items() if count >= min_count]
    return frequent_words

def load_text_from_docx(file_path):
    from docx import Document
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

def save_text_to_docx(file_path, text):
    from docx import Document
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    doc.save(file_path)

def get_document_info(file_path, min_count=2):
    text = load_text_from_docx(file_path)
    char_count = count_characters(text)
    frequent_words = count_frequent_words(text, min_count)
    return char_count, frequent_words
